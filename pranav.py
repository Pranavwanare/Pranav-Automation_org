import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import Table

# ---------------- SETTINGS -----------------
INPUT_ROOT = r"C:\Users\pranav.wanare\Downloads\OneDrive_2025-12-08\EOCs State-Wise"
OUTPUT_ROOT = r"C:\Users\pranav.wanare\Downloads\output"
TARGET_HEADER = "Cost Sharing Features"
# -------------------------------------------

# Formatting constants
INDENT_LEFT = Cm(0.64)
INDENT_RIGHT = Cm(0.82)
SPACE_BEFORE = Pt(4.9)
SPACE_AFTER = Pt(0)
BODY_FONT = "Arial"

HEADING_FONT = "Arial"
HEADING_SIZE = Pt(24)
SECTION_FONT = "Arial"
SECTION_SIZE = Pt(18)
BODY_SIZE = Pt(12)

def add_toc(paragraph):
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    paragraph._p.append(fld)

def copy_run_format(src_run, dst_run):
    dst_run.text = src_run.text
    dst_run.bold = src_run.bold
    dst_run.italic = src_run.italic
    dst_run.underline = src_run.underline
    dst_run.font.name = src_run.font.name or BODY_FONT
    dst_run.font.size = src_run.font.size or BODY_SIZE
    if src_run.font.color and src_run.font.color.rgb:
        dst_run.font.color.rgb = src_run.font.color.rgb
    else:
        dst_run.font.color.rgb = RGBColor(0, 0, 0)

def copy_paragraph(src_para, dst_doc):
    new_para = dst_doc.add_paragraph()
    
    # Determine style based on source
    style_name = src_para.style.name if src_para.style else ""
    if style_name.startswith("Heading"):
        new_para.alignment = 1  # center
        para_font_name = HEADING_FONT
        para_font_size = HEADING_SIZE
        para_bold = True
    elif style_name.lower().startswith("heading"):
        new_para.alignment = 0  # left
        para_font_name = SECTION_FONT
        para_font_size = SECTION_SIZE
        para_bold = True
    else:
        new_para.alignment = 0
        para_font_name = BODY_FONT
        para_font_size = BODY_SIZE
        para_bold = False

    new_para.paragraph_format.left_indent = INDENT_LEFT
    new_para.paragraph_format.right_indent = INDENT_RIGHT
    new_para.paragraph_format.space_before = SPACE_BEFORE
    new_para.paragraph_format.space_after = SPACE_AFTER
    new_para.paragraph_format.line_spacing = src_para.paragraph_format.line_spacing

    # Copy all runs
    for run in src_para.runs:
        new_run = new_para.add_run(run.text)
        copy_run_format(run, new_run)
        new_run.font.name = para_font_name
        new_run.font.size = para_font_size
        new_run.bold = para_bold or run.bold

def iter_block_items(doc):
    body = doc.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield child, Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield child, Table(child, doc)

def find_section_range(doc, target_header):
    target = target_header.strip().lower().replace(":", "")
    blocks = list(iter_block_items(doc))
    header_idx = None
    target_level = None

    for i, (elm, obj) in enumerate(blocks):
        if isinstance(obj, Paragraph):
            text = obj.text.strip().lower().replace(":", "")
            if text == target:
                header_idx = i
                style_name = obj.style.name if obj.style else ""
                if style_name.startswith("Heading"):
                    try:
                        target_level = int(style_name.replace("Heading", "").strip())
                    except:
                        target_level = 1
                else:
                    target_level = 1
                break

    if header_idx is None:
        return None, None

    end_idx = len(blocks) - 1
    for j in range(header_idx + 1, len(blocks)):
        elm, obj = blocks[j]
        if isinstance(obj, Paragraph):
            style_name = obj.style.name if obj.style else ""
            if style_name.startswith("Heading"):
                try:
                    level = int(style_name.replace("Heading", "").strip())
                except:
                    level = 1
                if level <= target_level:
                    end_idx = j - 1
                    break
    return header_idx, end_idx

def trim_document_to_section(doc, target_header):
    start_idx, end_idx = find_section_range(doc, target_header)
    if start_idx is None:
        return False

    body = doc.element.body
    blocks = list(body.iterchildren())

    # Preserve last section properties
    sectPr = None
    if blocks[-1].tag.endswith("sectPr"):
        sectPr = blocks[-1]
        blocks = blocks[:-1]

    # Remove blocks outside target
    for i in range(len(blocks) - 1, -1, -1):
        if i < start_idx or i > end_idx:
            body.remove(blocks[i])

    if sectPr is not None:
        body.append(sectPr)

    # Remove blank first paragraph
    if len(body) > 0:
        first_elem = body[0]
        if first_elem.tag.endswith("p") and not first_elem.text:
            body.remove(first_elem)

    return True

def create_output_doc(doc, output_path):
    # Add TOC at the top
    toc_para = doc.add_paragraph()
    add_toc(toc_para)
    doc.save(output_path)

def process_all():
    for root, dirs, files in os.walk(INPUT_ROOT):
        rel_path = os.path.relpath(root, INPUT_ROOT)
        out_dir = os.path.join(OUTPUT_ROOT, rel_path)
        os.makedirs(out_dir, exist_ok=True)

        for file in files:
            if not file.lower().endswith(".docx"):
                continue
            in_path = os.path.join(root, file)
            out_path = os.path.join(out_dir, file)

            print(f"Processing: {in_path}")
            try:
                doc = Document(in_path)
            except Exception as e:
                print(f"❌ Could not open {file}: {e}")
                continue

            found = trim_document_to_section(doc, TARGET_HEADER)
            if found:
                try:
                    create_output_doc(doc, out_path)
                    print(f"   ✔ Saved section to: {out_path}")
                except Exception as e:
                    print(f"   ❌ Could not save {file}: {e}")
            else:
                print(f"   ❌ Header '{TARGET_HEADER}' not found in: {file}")

    print("\n✔ All documents processed!")

if __name__ == "__main__":
    process_all()
